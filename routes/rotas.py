import os
import re
from io import BytesIO

from flask import (
    Blueprint,
    flash,
    jsonify,
    redirect,
    render_template,
    request,
    send_file,
    url_for,
)
from docx import Document
from werkzeug.utils import secure_filename
import pandas as pd

from forms import OrcamentoForm
from models import Alimento, db

rotas_blueprint = Blueprint('rotas', __name__)


@rotas_blueprint.route('/')
def index():
    """Display the dashboard with all items in stock."""
    alimentos = Alimento.query.all()
    return render_template('index.html', alimentos=alimentos)


@rotas_blueprint.route('/add', methods=['POST'])
def add():
    """Add a new item to the stock."""
    nome = request.form.get('nome', '').strip()
    quantidade_raw = request.form.get('quantidade', '')
    tipo_animal = request.form.get('tipo_animal', '').strip()
    reserva_raw = request.form.get('reserva', '')

    if not nome or not tipo_animal:
        flash('Nome e tipo de animal são obrigatórios.', 'error')
        return redirect(url_for('rotas.index'))

    try:
        quantidade = int(quantidade_raw)
        if quantidade < 0:
            flash('Quantidade não pode ser negativa.', 'error')
            return redirect(url_for('rotas.index'))
    except (ValueError, TypeError):
        flash('Quantidade deve ser um número válido.', 'error')
        return redirect(url_for('rotas.index'))

    reserva = None
    if reserva_raw:
        try:
            reserva = int(reserva_raw)
            if reserva < 0:
                flash('Reserva não pode ser negativa.', 'error')
                return redirect(url_for('rotas.index'))
        except (ValueError, TypeError):
            flash('Reserva deve ser um número válido.', 'error')
            return redirect(url_for('rotas.index'))

    try:
        alimento = Alimento(nome=nome, quantidade=quantidade, tipo_animal=tipo_animal, reserva=reserva)
        db.session.add(alimento)
        db.session.commit()
        flash('Item adicionado com sucesso!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Erro ao adicionar item: {str(e)}', 'error')

    return redirect(url_for('rotas.index'))


@rotas_blueprint.route('/delete/<int:id>', methods=['POST'])
def delete(id):
    """Delete an item from the stock."""
    alimento = Alimento.query.get_or_404(id)
    try:
        db.session.delete(alimento)
        db.session.commit()
        flash('Item removido com sucesso!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Erro ao remover item: {str(e)}', 'error')
    return redirect(url_for('rotas.index'))


@rotas_blueprint.route('/edit/<int:id>', methods=['GET', 'POST'])
def edit(id):
    """Edit an existing item in the stock."""
    alimento = Alimento.query.get_or_404(id)
    
    if request.method == 'POST':
        nome = request.form.get('nome', '').strip()
        quantidade_raw = request.form.get('quantidade', '')
        tipo_animal = request.form.get('tipo_animal', '').strip()
        reserva_raw = request.form.get('reserva', '')

        if not nome or not tipo_animal:
            flash('Nome e tipo de animal são obrigatórios.', 'error')
            return redirect(url_for('rotas.edit', id=id))

        try:
            quantidade = int(quantidade_raw)
            if quantidade < 0:
                flash('Quantidade não pode ser negativa.', 'error')
                return redirect(url_for('rotas.edit', id=id))
        except (ValueError, TypeError):
            flash('Quantidade deve ser um número válido.', 'error')
            return redirect(url_for('rotas.edit', id=id))

        reserva = None
        if reserva_raw:
            try:
                reserva = int(reserva_raw)
                if reserva < 0:
                    flash('Reserva não pode ser negativa.', 'error')
                    return redirect(url_for('rotas.edit', id=id))
            except (ValueError, TypeError):
                flash('Reserva deve ser um número válido.', 'error')
                return redirect(url_for('rotas.edit', id=id))

        try:
            alimento.nome = nome
            alimento.quantidade = quantidade
            alimento.tipo_animal = tipo_animal
            alimento.reserva = reserva
            db.session.commit()
            flash('Item atualizado com sucesso!', 'success')
            return redirect(url_for('rotas.index'))
        except Exception as e:
            db.session.rollback()
            flash(f'Erro ao atualizar item: {str(e)}', 'error')
    
    return render_template('edit.html', alimento=alimento)


@rotas_blueprint.route('/export')
def export():
    """Export stock data as CSV."""
    alimentos = Alimento.query.all()
    if not alimentos:
        flash('Não há itens para exportar.', 'warning')
        return redirect(url_for('rotas.index'))
    
    df = pd.DataFrame([a.to_dict() for a in alimentos])

    output = BytesIO()
    df.to_csv(output, index=False)
    output.seek(0)

    return send_file(output, mimetype='text/csv', as_attachment=True, download_name='estoque.csv')


@rotas_blueprint.route('/import', methods=['POST'])
def import_data():
    """Import stock data from CSV."""
    file = request.files.get('file')
    if not file or not file.filename or not file.filename.endswith('.csv'):
        flash('Por favor, selecione um arquivo CSV válido.', 'error')
        return redirect(url_for('rotas.index'))

    try:
        df = pd.read_csv(file)
        required_columns = ['nome', 'quantidade', 'tipo_animal']
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            flash(f'Colunas obrigatórias ausentes: {", ".join(missing_columns)}', 'error')
            return redirect(url_for('rotas.index'))
    except Exception as e:
        flash(f'Erro ao ler arquivo CSV: {str(e)}', 'error')
        return redirect(url_for('rotas.index'))

    try:
        imported_count = 0
        for _, row in df.iterrows():
            nome = str(row.get('nome', '')).strip()
            if not nome:
                continue
            try:
                quantidade = int(row.get('quantidade', 0))
                if quantidade < 0:
                    continue
            except (ValueError, TypeError):
                quantidade = 0
            tipo_animal = str(row.get('tipo_animal', '')).strip()
            if not tipo_animal:
                continue
                
            reserva = row.get('reserva')
            if reserva is not None:
                try:
                    reserva = int(reserva)
                    if reserva < 0:
                        reserva = None
                except (ValueError, TypeError):
                    reserva = None

            alimento = Alimento(
                nome=nome,
                quantidade=quantidade,
                tipo_animal=tipo_animal,
                reserva=reserva,
            )
            db.session.add(alimento)
            imported_count += 1

        db.session.commit()
        flash(f'{imported_count} itens importados com sucesso!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Erro ao importar dados: {str(e)}', 'error')
        
    return redirect(url_for('rotas.index'))


@rotas_blueprint.route('/orcamento', methods=['GET', 'POST'])
def orcamento():
    """Display and process budget form."""
    form = OrcamentoForm()
    if form.validate_on_submit():
        return redirect(url_for('rotas.gerar_orcamento'))
    return render_template('orcamento.html', form=form)


@rotas_blueprint.route('/gerar_orcamento', methods=['POST'])
def gerar_orcamento():
    """Generate a budget document in Word format."""
    nome_empresa = request.form.get('nome_empresa', '').strip()
    alimento = request.form.get('alimento', '').strip()
    try:
        valor_alimento = float(request.form.get('valor_alimento', 0))
        if valor_alimento < 0:
            flash('Valor não pode ser negativo.', 'error')
            return redirect(url_for('rotas.orcamento'))
    except (ValueError, TypeError):
        flash('Valor deve ser um número válido.', 'error')
        return redirect(url_for('rotas.orcamento'))

    if not nome_empresa or not alimento:
        flash('Nome da empresa e alimento são obrigatórios.', 'error')
        return redirect(url_for('rotas.orcamento'))

    try:
        doc = Document()
        doc.add_heading(f'Orçamento - {nome_empresa}', 0)
        doc.add_paragraph(f'Produto: {alimento}')
        doc.add_paragraph(f'Valor: R$ {valor_alimento:.2f}')

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        filename = re.sub(r'[^\w\-]', '_', nome_empresa)
        return send_file(output, as_attachment=True, download_name=f'orcamento_{filename}.docx')
    except Exception as e:
        flash(f'Erro ao gerar orçamento: {str(e)}', 'error')
        return redirect(url_for('rotas.orcamento'))


@rotas_blueprint.route('/generate_document')
def generate_document():
    """Generate a stock replenishment order document."""
    alimentos = Alimento.query.all()
    if not alimentos:
        flash('Não há itens em estoque para gerar pedido.', 'warning')
        return redirect(url_for('rotas.index'))
        
    try:
        document = Document()
        document.add_heading('Pedido de Novo Estoque', level=1)

        for alimento in alimentos:
            document.add_paragraph(f'{alimento.nome}: {alimento.quantidade} para {alimento.tipo_animal}')

        output = BytesIO()
        document.save(output)
        output.seek(0)

        return send_file(output, download_name='pedido_estoque.docx', as_attachment=True)
    except Exception as e:
        flash(f'Erro ao gerar documento: {str(e)}', 'error')
        return redirect(url_for('rotas.index'))


@rotas_blueprint.route('/create_document', methods=['POST'])
def create_document():
    """Generate a document from TinyMCE editor content."""
    content = request.form.get('tinymce', '').strip()
    if not content:
        flash('Conteúdo do documento não pode estar vazio.', 'error')
        return redirect(url_for('rotas.index'))
        
    try:
        doc = Document()
        doc.add_heading('Documento Editado', 0)

        clean_text = re.sub(r'<[^>]+>', '', content)
        doc.add_paragraph(clean_text)

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(output, as_attachment=True, download_name='documento_editado.docx')
    except Exception as e:
        flash(f'Erro ao criar documento: {str(e)}', 'error')
        return redirect(url_for('rotas.index'))


@rotas_blueprint.route('/api/alimentos', methods=['GET'])
def obter_alimentos():
    """API endpoint to get all items in JSON format."""
    alimentos = Alimento.query.all()
    return jsonify([a.to_dict() for a in alimentos])
