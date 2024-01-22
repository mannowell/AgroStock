import os
import json
from xml.dom.expatbuilder import DOCUMENT_NODE
from xml.dom.xmlbuilder import DocumentLS
from flask import send_from_directory
from xml.dom.minidom import DocumentFragment, DocumentType
from flask import jsonify, Flask, redirect, render_template, request, send_file
from wtforms import StringField  # ou outros campos que você estiver usando
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
from flask_wtf import FlaskForm
from flask_wtf.file import FileField, FileRequired
from wtforms import TextAreaField
from docx import Document
import pandas as pd
from werkzeug.utils import secure_filename
from io import BytesIO
import secrets

app = Flask(__name__)
data_folder = os.path.join(os.path.dirname(__file__), 'data')
os.makedirs(data_folder, exist_ok=True)

app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///' + os.path.join(data_folder, 'estoque.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SECRET_KEY'] = secrets.token_hex(16)
db = SQLAlchemy(app)
migrate = Migrate(app, db)

with app.app_context():
    db.create_all()
    
class DocumentForm(FlaskForm):
    documento = FileField('Documento', validators=[FileRequired()])
    
class OrcamentoForm(FlaskForm):
        # Defina os campos do formulário aqui
    nome_empresa = StringField('Nome da Empresa')
    alimento = StringField('Alimento')
    valor_alimento = StringField('Valor do Alimento')

class Alimento(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    nome = db.Column(db.String(50), nullable=False)
    quantidade = db.Column(db.Integer, nullable=False)
    tipo_animal = db.Column(db.String(20), nullable=False)
    reserva = db.Column(db.Integer)
   

@app.route('/')
def index():
    alimentos = Alimento.query.all()
    return render_template('index.html', alimentos=alimentos)

@app.route('/add', methods=['POST'])
def add():
    nome = request.form.get('nome')
    quantidade = int(request.form.get('quantidade'))
    tipo_animal = request.form.get('tipo_animal')

    alimento = Alimento(nome=nome, quantidade=quantidade, tipo_animal=tipo_animal)
    db.session.add(alimento)
    db.session.commit()

    return redirect('/')

@app.route('/export')
def export():
    alimentos = Alimento.query.all()
    df = pd.DataFrame([(a.id, a.nome, a.quantidade, a.tipo_animal) for a in alimentos],
                      columns=['ID', 'Nome', 'Quantidade', 'Tipo de Animal'])

    csv_data = df.to_csv(index=False).encode()

    return send_file(BytesIO(csv_data),
                     mimetype='text/csv',
                     as_attachment=True,
                     download_name='estoque.csv')
    

@app.route('/import', methods=['POST'])
def import_data():
   file = request.files['file']
   
   if file:
        filename = secure_filename(file.filename)
        file.save(filename)

        with open(filename, 'r') as f:
            # Lógica para processar o arquivo e salvar no banco de dados
            pass

        return redirect('/')
   
   return "Nenhum arquivo enviado", 400

@app.route('/orcamento', methods=['GET'])
def orcamento():
    form = OrcamentoForm()

    if request.method == 'POST' and form.validate_on_submit():
        # Lógica para processar os dados do formulário
        nome_empresa = form.nome_empresa.data
        alimento = form.alimento.data
        valor_alimento = form.valor_alimento.data

        # Agora você pode usar esses dados para gerar o documento PDF

    return render_template('orcamento.html', form=form)

@app.route('/gerar_orcamento', methods=['POST'])
def gerar_orcamento():
    # Lógica para processar os dados do formulário
    nome_empresa = request.form.get('nome_empresa')
    alimento = request.form.get('alimento')
    valor_alimento = float(request.form.get('valor_alimento'))

    # Adicione a lógica para gerar o documento com base nos dados do formulário
    documento = Document()
    # Adicione o conteúdo do documento conforme necessário
    # ...

    # Salve o documento temporariamente (ou em um local desejado)
    temp_file_path = 'orcamento_temp.docx'
    documento.save(temp_file_path)

    # Retorne o documento para download
    return send_file(temp_file_path, as_attachment=True, download_name='orcamento.docx')


#@app.route('/create_document', methods=['POST'])
def create_document():
    document = Document()

    # Adicione conteúdo ao documento com base nas informações fornecidas
    document.add_heading('ORÇAMENTO', level=1)

# Exemplo: Lendo o conteúdo de um arquivo
caminho_do_arquivo = 'C:\\Users\\Oliveiras\\Desktop\\Projeto\\arquivo.txt'
#try:
    #with open(caminho_do_arquivo, 'r') as arquivo:
           # conteudo = arquivo.read()
           # DocumentType.add_paragraph(conteudo)

    #return "O arquivo não foi encontrado", 404

    # Salve o documento temporariamente (ou em um local desejado)
    #temp_file_path = 'temp_document.docx'
 #   DocumentLS.save(temp_file_path)

    # Retorne o documento para download
    #return send_from_directory('.', 'temp_document.docx', as_attachment=True, download_name='orcamento.docx')


@app.route('/generate_document')
def generate_document():
    alimentos = Alimento.query.all()
    document = Document()
    document.add_heading('Pedido de Novo Estoque', level=1)

    for alimento in alimentos:
        document.add_paragraph(f'{alimento.nome}: {alimento.quantidade} para {alimento.tipo_animal}')

    output = BytesIO()
    document.save(output)
    output.seek(0)

    return send_file(output, download_name='pedido_estoque.docx', as_attachment=True)

    

@app.route('/generate_document_page', methods=['GET', 'POST'])
def generate_document_page():
    form = DocumentForm()

    if request.method == 'POST' and form.validate_on_submit():
        # Obtenha o documento do formulário
        documento = form.documento.data

        # Faça algo com o documento (por exemplo, gerar um arquivo Word)

        # Retorne o documento para download
        return send_file(documento, download_name='documento_editavel.docx', as_attachment=True)

    return render_template('generate_document_page.html', form=form)



@app.route('/edit_document')
def edit_document():
    # Pode ser necessário preencher o formulário com o conteúdo do documento existente
    # Se necessário, adicione lógica para carregar o conteúdo do documento na área de texto

    form = DocumentForm()
    
    initial_content = "Conteúdo inicial do documento"
    
    return render_template('edit_document.html', form=form)


@app.route('/api/alimentos', methods=['GET'])
def obter_alimentos():
    alimentos = Alimento.query.all()
     # Converte os objetos SQLAlchemy para um formato que pode ser serializado para JSON
    alimentos_serializados = [
        {
            'id': alimento.id,
            'nome': alimento.nome,
            'quantidade': alimento.quantidade,
            'tipo_animal': alimento.tipo_animal,
            # Adicione mais campos conforme necessário
        }
        for alimento in alimentos
    ]

    return jsonify(alimentos_serializados)


if __name__ == '__main__':
    app.run(debug=True)
